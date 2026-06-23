# -*- coding: utf-8 -*-
"""Adapt the Mattias Gustawsson konsultavtal into a CEO template (Swiss company, placeholders)."""
from docx import Document
import copy

SRC="/root/.claude/uploads/cc747475-841f-5d65-b8c7-9a970e1dee23/33e696b7-Mattias_Gustawsson_avtal.docx"
OUT="/home/user/Claude/Konsultavtal_CEO_template.docx"
d=Document(SRC)
P=d.paragraphs

def set_text(p,text):
    """Replace the whole paragraph text with one run, preserving the first run's size/bold/italic."""
    size=p.runs[0].font.size if p.runs else None
    name=p.runs[0].font.name if p.runs else None
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run=p.add_run(text)
    if size is not None: run.font.size=size
    if name is not None: run.font.name=name
    return p

# --- 9: counterparty = Swiss company, services performed by CEO ---
set_text(P[9],
 "[Bolagsnamn], ett bolag registrerat i Schweiz, med adress [ange adress], nedan kallat ”Konsulten”. "
 "Tjänsterna ska utföras av [Konsultens namn] i egenskap av verkställande direktör (CEO) för Uppdragsgivaren.")

# --- 12: Uppdraget (CEO duties) ---
set_text(P[12],
 "Konsulten åtar sig att som självständig konsult tillhandahålla tjänster åt Uppdragsgivaren genom att "
 "[Konsultens namn] verkar som verkställande direktör (CEO) för Uppdragsgivaren. I rollen ingår att leda och "
 "ansvara för bolagets löpande verksamhet och förvaltning i enlighet med styrelsens och Uppdragsgivarens "
 "instruktioner samt gällande rätt.")

# --- 14: Avtalstid (dates -> placeholders) ---
set_text(P[14],
 "Detta avtal träder i kraft den [startdatum] och upphör automatiskt den [slutdatum]. Båda parter, "
 "Uppdragsgivaren och Konsulten, har rätt att avsluta uppdraget omedelbart om vissa kriterier uppfylls, såsom om "
 "verksamhetens behov förändras, vid bristande prestation, eller om projektet avslutas i förtid. Om kontraktet "
 "avslutas i förtid sker det utan ytterligare kostnader.")

# --- 16: Ersättning heading stays; 17: fixed fee in brackets ---
set_text(P[17],
 "För utförandet av uppdraget erhåller Konsulten en fast ersättning om [CHF 21 500 per månad].")

# --- insert performance-units paragraph as a separate paragraph after the fixed fee ---
units_text=(
 "Utöver den fasta ersättningen tilldelas Konsulten 8 prestationsenheter (performance units) i enlighet med "
 "Uppdragsgivarens vid var tid gällande ersättningssystem, ”Roles, Responsibilities & Performance Compensation "
 "Framework”. Prestationsenheterna är giltiga endast så länge [Konsultens namn] personligen arbetar för och utför "
 "tjänsterna åt Uppdragsgivaren. När [Konsultens namn] upphör att arbeta för Uppdragsgivaren upphör samtliga "
 "prestationsenheter och därtill hörande rättigheter att gälla med omedelbar verkan, oavsett om [Konsultens namn] "
 "anses vara en good leaver eller bad leaver. Prestationsenheterna kan inte överlåtas.")
newp=P[18].insert_paragraph_before("")          # insert between 17 and the blank line at 18
newp.style=P[17].style
r=newp.add_run(units_text)
if P[17].runs and P[17].runs[0].font.size is not None: r.font.size=P[17].runs[0].font.size

# --- insert "Personligt åtagande" clause after Uppdraget (before Avtalstid at P[13]) ---
SZ=P[12].runs[0].font.size if P[12].runs else None
ph=P[13].insert_paragraph_before("")             # heading
ph.style=P[16].style                             # match "Body Text 2" section heading
hr=ph.add_run("Personligt åtagande"); hr.font.bold=True
if SZ is not None: hr.font.size=SZ
pb=P[13].insert_paragraph_before("")             # body (goes after heading, before Avtalstid)
pb.style=P[14].style                             # "Body Text"
br=pb.add_run(
 "Tjänsterna är av personlig natur och ska utföras personligen av [Konsultens namn]. [Konsultens namn] gör genom "
 "detta avtal ett personligt åtagande gentemot Uppdragsgivaren att utföra uppdraget och att iaktta de skyldigheter "
 "som följer av avtalet, inklusive sekretess och de villkor som gäller för prestationsenheterna. Konsulten får inte "
 "ersätta [Konsultens namn] med annan person eller överlåta utförandet av tjänsterna utan Uppdragsgivarens "
 "skriftliga medgivande. [Konsultens namn] undertecknar detta avtal såväl för Konsultens räkning som personligen "
 "för att bekräfta detta åtagande.")
if SZ is not None: br.font.size=SZ

# --- 19: invoicing (Konsulten; specify fixed fee) ---
set_text(P[19],
 "Konsulten ska fakturera Uppdragsgivaren månadsvis i efterskott. Betalningsvillkor är 30 dagar netto från "
 "fakturadatum. Fakturan ska specificera den fasta ersättningen samt, i förekommande fall, prestationsbaserad "
 "ersättning enligt ersättningssystemet.")

# --- 21: relationship between parties (company, not individual) ---
set_text(P[21],
 "Konsulten är ett självständigt bolag och detta avtal medför inte ett anställningsförhållande. Konsulten ansvarar "
 "själv för skatter, sociala avgifter, försäkringar m.m.")

# --- 23: confidentiality (MG -> Konsulten; honom -> Konsulten; anställning -> uppdrag) ---
set_text(P[23],
 "Konsulten får inte utan Bolagets uttryckliga medgivande för tredje man avslöja sådant som angår Bolagets, dess "
 "dotter- eller intressebolags eller samarbetspartners verksamhet eller affärsangelägenheter i vidare mån än vad som "
 "krävs för uppdragets behöriga utförande. Konsulten förbinder sig att ej heller för egen eller annans del begagna "
 "sig av vad som i detta hänseende blivit Konsulten bekant. Detta sekretessåtagande gäller även efter uppdragets "
 "upphörande.")

# --- 25: return of documents ---
set_text(P[25],
 "Vid uppdragets upphörande åligger det Konsulten att till Bolaget återlämna alla affärshandlingar av vad slag det "
 "vara må rörande Bolagets, dess dotter- eller intressebolags eller samarbetspartners angelägenheter som Konsulten "
 "har i sin besittning eller som denne på annat sätt har tillgång till.")

# --- 28: guidelines ---
set_text(P[28],
 "Konsulten ska vid utförandet av uppdraget iaktta Uppdragsgivarens anvisningar samt gällande rätt. Därutöver har "
 "Konsulten att vid var tid bevaka och tillvarata Uppdragsgivarens intressen samt fullgöra sina uppgifter efter "
 "bästa förmåga.")

# --- 30: entire agreement ---
set_text(P[30],
 "Detta avtal utgör hela överenskommelsen mellan Uppdragsgivaren och Konsulten. Inga andra muntliga eller skriftliga "
 "överenskommelser är giltiga förutom de interna policies, dvs. de policies som gäller för anställda hos "
 "Uppdragsgivaren, vilka även gäller Konsulten i sin roll som konsult.")

# --- 34: arbitration costs ---
set_text(P[34],
 "Uppdragsgivaren skall svara för samtliga ombuds- och skiljemannakostnader med undantag för två basbelopp enligt "
 "lagen för allmän försäkring som skall betalas av Konsulten. Med basbelopp avses härmed det basbelopp som gäller "
 "vid tidpunkten för inledande av skiljeförfarandet. Om skiljemännen finner att Konsulten skäligen föranlett "
 "skiljeförfarandet, skall skiljemännen äga rätt att fördela kostnaderna efter eget skön.")

# --- 37: Övrigt (attestation about the person performing services) ---
set_text(P[37],
 "Konsulten intygar i och med detta avtal att den person som utför tjänsterna inte är dömd i domstol och inte har "
 "varit föremål för polisutredning eller liknande tidigare. Konsulten förbinder sig också att meddela "
 "Uppdragsgivaren om sådan situation skulle uppstå under uppdragstiden.")

# --- 44: place and date placeholder ---
set_text(P[44],"[Ort] den [datum]")

# --- 51: signature line (replace Mattias, keep tabs) + capacity line ---
for r in P[51].runs:
    if "Mattias Gustawsson" in r.text:
        r.text=r.text.replace("Mattias Gustawsson","[Konsultens namn]")
cap=d.add_paragraph()                            # capacity labels under the names
cap.style=P[51].style
cr=cap.add_run("för Uppdragsgivaren\t\t\tför [Bolagsnamn] och personligen")
if P[51].runs and P[51].runs[0].font.size is not None: cr.font.size=P[51].runs[0].font.size

d.save(OUT)
print("Saved",OUT)
# verify
d2=Document(OUT)
for i,p in enumerate(d2.paragraphs):
    if p.text.strip(): print(f"[{i}] ({p.style.name}) {p.text}")
