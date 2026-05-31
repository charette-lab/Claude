"""Research-standard references & methodology page for the Athanase-vs-Private-
Equity analysis. Full bibliographic citations with DOIs (hanging indent),
rendered as a brand-styled PDF document.

Academic citations use verified DOIs (several supplied by the client).
Practitioner / dataset sources are cited by institution, title and year with a
URL where no formal DOI exists — and are labelled as practitioner research so
the distinction from peer-reviewed work is explicit.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

NAVY = RGBColor(0x15, 0x21, 0x30)
SLATE = RGBColor(0x31, 0x43, 0x59)
BODY = RGBColor(0x0E, 0x16, 0x20)
GREY = RGBColor(0x55, 0x6A, 0x83)

doc = Document()
# Letter portrait, research-paper margins
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.9)
sec.left_margin = sec.right_margin = Inches(1.0)

base = doc.styles["Normal"]
base.font.name = "Times New Roman"
base.font.size = Pt(10.5)
base.font.color.rgb = BODY


def _set(run, name="Times New Roman", size=10.5, color=BODY, bold=False,
         italic=False):
    run.font.name = name; run.font.size = Pt(size)
    run.font.color.rgb = color; run.font.bold = bold; run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def heading(text, size=13, space_before=14, space_after=6, color=SLATE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    _set(p.add_run(text), size=size, color=color, bold=True)
    return p


def reference(parts):
    """parts: list of (text, italic) tuples; rendered with hanging indent."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.3)
    pf.first_line_indent = Inches(-0.3)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.05
    for text, italic in parts:
        _set(p.add_run(text), size=10, italic=italic)
    return p


def note(text, size=9.5, italic=True, color=GREY, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    _set(p.add_run(text), size=size, italic=italic, color=color)
    return p


# ---- Title block -----------------------------------------------------------
t = doc.add_paragraph()
t.paragraph_format.space_after = Pt(2)
_set(t.add_run("Athanase Industrial Partner"), size=11, color=GREY, bold=True)
h = doc.add_paragraph()
h.paragraph_format.space_after = Pt(2)
_set(h.add_run("Private Equity Comparison — References & Methodology"),
     size=17, color=NAVY, bold=True)
sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(8)
_set(sub.add_run("Full source list for the risk, return and correlation "
                 "figures used in the Athanase-vs-private-equity analysis. "
                 "Peer-reviewed literature is cited with DOIs; practitioner and "
                 "dataset sources are identified by institution and year."),
     size=10, italic=True, color=GREY)

# divider rule (bottom border on an empty paragraph)
from docx.oxml import OxmlElement
rule = doc.add_paragraph()
rule.paragraph_format.space_after = Pt(2)
_pPr = rule._p.get_or_add_pPr()
_pbdr = OxmlElement("w:pBdr")
_bottom = OxmlElement("w:bottom")
_bottom.set(qn("w:val"), "single")
_bottom.set(qn("w:sz"), "6")
_bottom.set(qn("w:space"), "1")
_bottom.set(qn("w:color"), "314359")
_pbdr.append(_bottom)
_pPr.append(_pbdr)

# ===========================================================================
# 1. Peer-reviewed and academic literature
# ===========================================================================
heading("1.  Peer-reviewed and academic literature", space_before=8)

reference([("Ang, A., Chen, B., Goetzmann, W. N., & Phalippou, L. (2018). "
            "Estimating private equity returns from limited partner cash flows. ",
            False),
           ("The Journal of Finance, 73", True),
           ("(4), 1751–1783. https://doi.org/10.1111/jofi.12688", False)])

reference([("Boyer, B. H., Nadauld, T. D., Vorkink, K. P., & Weisbach, M. S. "
            "(2018). Private equity indices based on secondary market "
            "transactions. ", False),
           ("SSRN Electronic Journal", True),
           (" (also NBER Working Paper No. 25207). "
            "https://doi.org/10.2139/ssrn.3272357", False)])

reference([("Buchner, A., Kaserer, C., & Wagner, N. F. (2010). Private equity "
            "funds: Valuation, systematic risk and illiquidity. ", False),
           ("SSRN Electronic Journal", True),
           (". https://doi.org/10.2139/ssrn.1102471", False)])

reference([("Couts, S., Gonçalves, A. S., & Rossi, A. (2020). Unsmoothing "
            "returns of illiquid funds. ", False),
           ("SSRN Electronic Journal", True),
           (". https://doi.org/10.2139/ssrn.3544854", False)])

reference([("Getmansky, M., Lo, A. W., & Makarov, I. (2004). An econometric "
            "model of serial correlation and illiquidity in hedge fund "
            "returns. ", False),
           ("Journal of Financial Economics, 74", True),
           ("(3), 529–609. https://doi.org/10.1016/j.jfineco.2004.04.001",
            False)])

reference([("Harris, R. S., Jenkinson, T., & Kaplan, S. N. (2014). Private "
            "equity performance: What do we know? ", False),
           ("The Journal of Finance, 69", True),
           ("(5), 1851–1882. https://doi.org/10.1111/jofi.12154", False)])

reference([("Hayley, S., & Sefiloglu, O. (2022). Biases in private equity "
            "returns. ", False),
           ("SSRN Electronic Journal", True),
           (". https://doi.org/10.2139/ssrn.4245715", False)])

reference([("Jackson, B., Ling, D. C., & Naranjo, A. (2022). Catering and "
            "return manipulation in private equity. ", False),
           ("SSRN Electronic Journal", True),
           (". https://doi.org/10.2139/ssrn.4244467", False)])

reference([("Kieffer, E., Meyer, T., Gloukoviezoff, G., Lucius, H., & Bouvry, "
            "P. (2023). Learning private equity recommitment strategies for "
            "institutional investors. ", False),
           ("Frontiers in Artificial Intelligence, 6", True),
           (", 1014317. https://doi.org/10.3389/frai.2023.1014317", False)])

reference([("Meyer, T. (2020). Hidden in plain sight — the impact of undrawn "
            "commitments. ", False),
           ("The Journal of Alternative Investments, 23", True),
           ("(2), 94–110. https://doi.org/10.3905/jai.2020.1.101", False)])

reference([("Phalippou, L. (2020). An inconvenient fact: Private equity "
            "returns and the billionaire factory. ", False),
           ("The Journal of Investing, 30", True),
           ("(1), 11–39 (University of Oxford, Saïd Business School Working "
            "Paper 2020-10). https://doi.org/10.3905/joi.2020.1.153", False)])

# ===========================================================================
# 2. Industry datasets and practitioner research
# ===========================================================================
heading("2.  Industry datasets and practitioner research")
note("Practitioner reports are cited by institution and year; where no formal "
     "publication record or DOI exists, the issuing body and report series are "
     "given. Figures are as reported by the issuer.", after=6)

reference([("Asness, C. S. (2023). ", False),
           ("Why does private equity get to keep calling recessions "
            "investment opportunities? (“Volatility laundering”). ", True),
           ("AQR Capital Management — Cliff’s Perspectives.", False)])

reference([("British Private Equity & Venture Capital Association (BVCA) & "
            "Capital Dynamics. (2024). ", False),
           ("BVCA Private Equity and Venture Capital Performance Measurement "
            "Survey (PME+ methodology).", True),
           (" London: BVCA.", False)])

reference([("Cambridge Associates. (2024). ", False),
           ("US Private Equity Index and Selected Benchmark Statistics "
            "(Q2 2024).", True),
           (" Boston: Cambridge Associates LLC.", False)])

reference([("Cliffwater LLC. (2023). ", False),
           ("Long-term private equity performance in US state pension plans "
            "(2000–2022).", True),
           (" Marina del Rey: Cliffwater LLC.", False)])

reference([("Morningstar. (2023). ", False),
           ("Volatility laundering: How private equity funds understate the "
            "risk of their investments.", True),
           (" Chicago: Morningstar, Inc.", False)])

reference([("PIMCO. (2022). ", False),
           ("The discreet charm of private assets: De-smoothing and the true "
            "volatility of private equity.", True),
           (" Newport Beach: PIMCO.", False)])

reference([("Two Sigma Investments. (2024). ", False),
           ("The alternative truth of private equity (Venn by Two Sigma): "
            "de-smoothing the Preqin Private Equity Index.", True),
           (" New York: Two Sigma.", False)])

# ===========================================================================
# 3. Data and methodology
# ===========================================================================
heading("3.  Data and methodology")
note("All Athanase figures are computed directly from the fund’s audited net "
     "monthly return series, 235 months (2006–2025), not drawn from third-party "
     "estimates:", italic=False, color=BODY, after=4)

for txt in [
    "Annualised return: geometric compounding of monthly net returns; "
    "16.0% p.a.",
    "Total volatility: standard deviation of monthly returns × √12; 27.0% p.a.",
    "Downside volatility: deviation of negative months below a 0% minimum "
    "acceptable return × √12; 10.9% p.a.",
    "Correlation / beta: versus the MSCI World IMI over the common period; "
    "0.44 correlation, 0.73 beta.",
    "Large-loss frequency: share of rolling 36-month windows with a cumulative "
    "return below −30%; ≈2%.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3)
    _set(p.add_run(txt), size=9.5)

note("Private-equity figures are taken from the sources in Sections 1–2. "
     "Where a chart shows a “de-smoothed”, “true economic” or “committed-"
     "capital” value, it reflects the cited research rather than reported "
     "(appraisal-based) data. Bridge and illustrative magnitudes reconcile the "
     "cited endpoints and are labelled as such on the relevant slides. Past "
     "performance is not indicative of future results; this document is for "
     "professional investors and is not investment advice.",
     italic=True, color=GREY, after=2)

out = "Athanase_vs_PE_References.docx"
doc.save(out)
print("Saved", out)
