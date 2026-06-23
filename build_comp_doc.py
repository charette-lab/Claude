"""Staff-facing Roles, Responsibilities & Performance Compensation Framework (.docx) v2."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x54,0x96); GREY=RGBColor(0x59,0x59,0x59)
OUT="/home/user/Claude/Performance_Compensation_Framework.docx"

doc=Document()
sec=doc.sections[0]
sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.6); sec.left_margin=Cm(2.0); sec.right_margin=Cm(2.0)
st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(10.3)
st.paragraph_format.space_after=Pt(5); st.paragraph_format.line_spacing=1.05

def shade(cell,hexcolor):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexcolor); tcPr.append(sh)
def setcellfont(cell,size=9,bold=False,color=None,align=None):
    for p in cell.paragraphs:
        if align is not None: p.alignment=align
        p.paragraph_format.space_after=Pt(1); p.paragraph_format.space_before=Pt(1)
        for r in p.runs:
            r.font.size=Pt(size); r.font.bold=bold
            if color is not None: r.font.color.rgb=color
def heading(text,size=12.5,color=NAVY,space_before=8):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(space_before); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(text); r.font.bold=True; r.font.size=Pt(size); r.font.color.rgb=color; return p
def subhead(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(text); r.font.bold=True; r.font.size=Pt(10.6); r.font.color.rgb=BLUE; return p
def rich(segments,style_after=5,indent=None):
    """segments: list of (text, bold) tuples -> one paragraph."""
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(style_after)
    if indent: p.paragraph_format.left_indent=Pt(indent)
    for text,bold in segments:
        r=p.add_run(text); r.font.size=Pt(10.3); r.font.bold=bold
    return p
def body(text,after=5): return rich([(text,False)],after)
def bullet(segments):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(2)
    for text,bold in segments:
        r=p.add_run(text); r.font.size=Pt(10.3); r.font.bold=bold
    return p

# ---------- title ----------
t=doc.add_paragraph(); r=t.add_run("Athanase Industrial Partner"); r.font.bold=True; r.font.size=Pt(17); r.font.color.rgb=NAVY
t.paragraph_format.space_after=Pt(0)
s=doc.add_paragraph(); r=s.add_run("Roles, Responsibilities & Performance Compensation Framework"); r.font.size=Pt(12.5); r.font.color.rgb=BLUE
s.paragraph_format.space_after=Pt(2)
m=doc.add_paragraph(); r=m.add_run("Staff guide to how we work and how the performance pool is created and shared"); r.italic=True; r.font.size=Pt(10); r.font.color.rgb=GREY
m.paragraph_format.space_after=Pt(2)
pb=doc.add_paragraph(); pPr=pb._p.get_or_add_pPr(); bdr=OxmlElement("w:pBdr"); bot=OxmlElement("w:bottom")
bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"12"); bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),"1F3864")
bdr.append(bot); pPr.append(bdr); pb.paragraph_format.space_after=Pt(6)

# ---------- 1. purpose ----------
heading("1.  Purpose",space_before=2)
body("This document explains how we work at Athanase Industrial Partner — the roles in the firm and what each is "
     "responsible for — and how performance-based compensation is created and shared. Our philosophy is simple: when "
     "the fund performs, the people who create that performance share directly in it. The framework is designed to be "
     "transparent, predictable, and to reward both contribution and progression.")

# ---------- 2. roles & responsibilities ----------
heading("2.  Roles and responsibilities")
subhead("2.1  How we invest")
body("Athanase is an engaged owner of public companies. We run a concentrated portfolio and take meaningful ownership "
     "positions, then work actively to increase the value of each holding — through constructive engagement with "
     "boards and management on strategy, operations, capital allocation and governance. This makes our model closer "
     "to private equity than to a traditional trading desk: we create returns by improving businesses we own, not by "
     "trading in and out of them.")
body("Every prospective investment is owned by a deal team made up of a Portfolio Manager and at least one Analyst. "
     "The deal team develops the thesis and the value-creation plan and presents it to the wider investment group. "
     "The group's role is to pressure-test the investment premise — challenging assumptions, surfacing risks and "
     "stress-testing the thesis from every angle. The CIO chairs these meetings, makes the final selection decision, "
     "and owns portfolio construction — position sizing and the pace of accumulation and divestment — working closely "
     "with the responsible Portfolio Manager and the Head of Research.")

subhead("2.2  The investment team")
roles_long=[
("Chief Investment Officer (CIO). ","Holds ultimate responsibility for the firm's investment performance. The CIO chairs the investment meetings and makes the final decision on which investments enter and leave the portfolio. The CIO owns portfolio construction — determining position sizes and the pace of accumulation and divestment — in close collaboration with the responsible Portfolio Manager and the Head of Research. The CIO sets the investment strategy and safeguards the discipline and consistency of our process."),
("Portfolio Manager (PM). ","Leads assigned investments. Unlike a traditional hedge fund, a PM at Athanase is not allocated a discrete pool of capital to trade. As an engaged owner, the PM acts much like a private equity deal partner: they take ownership of each assigned investment and are responsible for driving the improvement of the underlying portfolio company. This includes building and maintaining the thesis and value-creation plan, engaging with the company's board and management, pursuing operational, strategic, capital-allocation and governance change, and representing the firm in its ownership role. The PM leads the deal team and presents and defends the thesis to the investment group."),
("Junior PM. ","Supports and co-leads investments alongside a Portfolio Manager, taking ownership of defined workstreams within the engagement and value-creation agenda. The role bridges deep analytical work and full investment ownership, building the judgment and company-engagement experience needed to progress to PM."),
("Head of Research. ","A distinct senior role that leads the firm's research function and partners closely with the CIO on portfolio construction — helping shape conviction, position sizing and risk assessment across the portfolio — while overseeing the rigour and quality of our due diligence. This role is currently held on a voluntary basis and does not participate in the performance pool."),
("Senior Analyst. ","A lead analyst with deep sector or thematic coverage who generates original investment ideas, leads the due-diligence effort on assigned investments and mentors junior analysts. The senior analyst is the most experienced member of the research bench and a primary feeder to the Junior PM and PM roles."),
("Analyst. ","Responsible for the confirmatory due diligence that underpins each investment decision and for helping to source new ideas. Working within a deal team, the analyst builds the financial models, gathers and verifies evidence, tests the key premises of the thesis and monitors investments once owned. Strong analysts are active idea generators and form the bench from which Senior Analysts and PMs are developed."),
]
for lead,rest in roles_long:
    rich([(lead,True),(rest,False)],style_after=4,indent=6)

subhead("2.3  Leadership, control and support")
roles_short=[
("Chief Executive Officer (CEO)"," — leads the firm as a business: strategy, commercial development, governance and overall management, so the investment team can focus on investing."),
("Chief Financial Officer (CFO)"," — owns the financial management of the firm and the funds: fund accounting, treasury, financial control, audit and tax."),
("Chief Operating Officer (COO)"," — responsible for the operating platform: operations, infrastructure, technology, vendors and the day-to-day running of the business."),
("Execution"," — implements the CIO's accumulation and divestment decisions efficiently and discreetly, given our concentrated ownership positions."),
("Risk"," — independent monitoring and management of portfolio and firm-wide risk, including exposure, liquidity and concentration."),
("Compliance"," — ensures the firm meets its regulatory obligations and upholds the highest standards of conduct."),
("Investor Relations (IR)"," — owns the investor relationship: reporting, communication and capital raising."),
("Support"," — provides the operational and administrative backbone that keeps the firm running."),
]
for lead,rest in roles_short:
    bullet([(lead,True),(rest,False)])

# ---------- 3. pool created ----------
heading("3.  How the performance pool is created")
rich([("When the fund delivers gains for our investors, it earns a performance fee. A fixed share of that performance fee — the ",False),
      ("team take of 60%",True),
      (" — is set aside to form the performance pool for staff. The remaining 40% is retained by the firm. The pool is "
       "therefore directly linked to the returns we generate: a stronger year for the fund means a larger pool, and as "
       "assets under management (AUM) grow, the pool grows with them.",False)])
rich([("The performance fee earned on the firm's ",False),("S class of shares",True),
      (" is excluded from the pool: performance units do not receive any share of the performance fee generated by the "
       "S class of shares.",False)])

# ---------- 4. weights & units ----------
heading("4.  How the pool is shared — performance weights and units")
body("Each role carries a performance weight that reflects its seniority and its contribution to investment "
     "performance. Weights convert the pool into individual awards through a simple, fully transparent calculation:")
for seg in [
    [("Each person's weight represents a number of ",False),("performance units",True),(".",False)],
    [("We add up the weights of everyone in the firm to get the ",False),("total units",True),(".",False)],
    [("The pool is divided by the total units to give the ",False),("value of one unit",True),(".",False)],
    [("Your award = ",False),("your weight × the value of one unit",True),(".",False)]]:
    bullet(seg)
fb=doc.add_table(rows=1,cols=1); fb.alignment=WD_TABLE_ALIGNMENT.CENTER
cell=fb.cell(0,0); shade(cell,"EAF0FA"); cell.width=Inches(6.2)
for line in ["Value of one unit  =  Performance pool  ÷  Total performance units",
             "Your performance award  =  Your weight  ×  Value of one unit"]:
    p=(cell.paragraphs[0] if line.startswith("Value") else cell.add_paragraph())
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(line); r.font.bold=True; r.font.size=Pt(10.5); r.font.color.rgb=NAVY
for bd in ("top","bottom","left","right"):
    e=OxmlElement(f"w:{bd}"); e.set(qn("w:val"),"single"); e.set(qn("w:sz"),"6"); e.set(qn("w:color"),"2E5496")
    cell._tc.get_or_add_tcPr().append(e)
doc.add_paragraph().paragraph_format.space_after=Pt(2)

# ---------- 5. roles & weights table ----------
heading("5.  Roles and performance weights")
body("The table below sets out the roles in the firm, grouped by track, and the performance weight attached to each. "
     "Weights are reviewed periodically and apply per person in the role.",after=4)
rows=[
 ("Investment","Chief Investment Officer (CIO)","Selection, chairs investment meetings, owns portfolio construction.","16"),
 ("Investment","Portfolio Manager (PM)","Leads investments; drives value creation in portfolio companies.","8"),
 ("Investment","Junior PM","Co-leads investments; owns workstreams of the engagement agenda.","6"),
 ("Investment","Head of Research","Leads research function; partners with CIO on portfolio construction.","—"),
 ("Investment","Senior Analyst","Original idea generation, leads due diligence, mentors analysts.","4"),
 ("Investment","Analyst","Confirmatory due diligence and idea sourcing within deal teams.","2.5"),
 ("Leadership","Chief Executive Officer (CEO)","Firm leadership and overall business management.","8"),
 ("Leadership","Chief Financial Officer (CFO)","Finance, fund accounting and treasury.","3"),
 ("Leadership","Chief Operating Officer (COO)","Operations and firm infrastructure.","3"),
 ("Control","Execution","Trade execution of accumulation and divestment.","2"),
 ("Control","Risk","Risk management and portfolio monitoring.","2"),
 ("Control","Compliance","Regulatory compliance.","1"),
 ("Client & Support","Investor Relations (IR)","Investor relations and capital raising.","1"),
 ("Client & Support","Support","Operational and administrative support.","1"),
]
tbl=doc.add_table(rows=1,cols=4); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.style="Table Grid"
hdr=tbl.rows[0].cells
for i,txt in enumerate(["Track","Role","Primary responsibility","Weight"]):
    hdr[i].text=txt; shade(hdr[i],"1F3864")
    setcellfont(hdr[i],size=9,bold=True,color=RGBColor(0xFF,0xFF,0xFF),
                align=WD_ALIGN_PARAGRAPH.CENTER if i==3 else WD_ALIGN_PARAGRAPH.LEFT)
track_fill={"Investment":"DDE9F7","Leadership":"E7EFDC","Control":"FBE9D8","Client & Support":"F2F2F2"}
for tr,role,desc,wt in rows:
    c=tbl.add_row().cells
    c[0].text=tr; c[1].text=role; c[2].text=desc; c[3].text=wt
    shade(c[0],track_fill[tr])
    setcellfont(c[0],size=8.5,bold=True,color=NAVY)
    setcellfont(c[1],size=8.5,bold=True)
    setcellfont(c[2],size=8.5)
    setcellfont(c[3],size=10,bold=True,color=BLUE,align=WD_ALIGN_PARAGRAPH.CENTER)
for row in tbl.rows:
    row.cells[0].width=Inches(1.1); row.cells[1].width=Inches(2.0); row.cells[2].width=Inches(2.9); row.cells[3].width=Inches(0.7)
fn=doc.add_paragraph(); fn.paragraph_format.space_before=Pt(2)
r=fn.add_run("“—”  The Head of Research is currently held on a voluntary basis and does not participate in the performance pool.")
r.italic=True; r.font.size=Pt(8); r.font.color.rgb=GREY

# ---------- 6. growth ----------
heading("6.  How awards behave as the firm grows")
body("Two things happen as the firm adds people and assets, and it is important to understand both:")
bullet([("Your ",False),("percentage share",True),(" of the pool may gradually reduce as the team grows and total units increase.",False)])
bullet([("Your ",False),("award in absolute terms is expected to rise",True),(", because the pool grows with AUM and fund performance.",False)])
body("In other words, a smaller slice of a much larger pie. This keeps the framework fair as we scale: building a "
     "stronger, larger team is what drives the bigger pool that benefits everyone.",after=4)

# ---------- 7. progression ----------
heading("7.  Career progression")
rich([("Weights increase with responsibility. The investment track in particular offers a clear ladder — ",False),
      ("Analyst → Senior Analyst → Junior PM → Portfolio Manager → CIO",True),
      (" — with each step carrying a meaningfully higher weight. Progression is how your share of the pool grows "
       "alongside the firm.",False)])

# ---------- 8. units and leaving ----------
heading("8.  Performance units and leaving the firm")
rich([("Performance units are the firm's ",False),("performance-related bonus mechanism",True),
      (": they are funded entirely from the firm's performance fees and represent a share of that bonus pool. They "
       "are not equity, ownership, or a vested or transferable asset.",False)])
rich([("Performance units are valid only for as long as you are working at the firm. When you leave — for any reason, "
       "and regardless of whether you are treated as a ",False),
      ("good leaver or a bad leaver",True),
      (" — your performance units are ",False),("relinquished in full",True),
      (" and lapse with immediate effect, and no further award is payable in respect of them. Units cannot be "
       "transferred or assigned.",False)])

# ---------- 9. governance ----------
heading("9.  Governance")
body("The framework, the list of roles and the weights are reviewed periodically by the firm's leadership. Awards are "
     "determined by this framework; the firm retains reasonable discretion in exceptional circumstances and in line "
     "with regulatory requirements. This document is provided for information and does not form part of any "
     "individual's contract of employment.",after=2)

foot=doc.add_paragraph()
r=foot.add_run("Athanase Industrial Partner  ·  Roles, Responsibilities & Performance Compensation Framework  ·  Internal — for staff information")
r.italic=True; r.font.size=Pt(8); r.font.color.rgb=GREY

doc.save(OUT)
print("Saved",OUT)
